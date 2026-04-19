import os
import pandas as pd
import sqlite3

import fitz
import docx
from docx import Document
import whisper # Для аудио формата
import time
import datetime
import xml.etree.ElementTree as ET
import io
import tempfile
import subprocess
import pytesseract
import polars
from PIL import Image
from concurrent.futures import ThreadPoolExecutor
import cv2
import numpy as np
import re
import mediapipe as mp
from striprtf.striprtf import rtf_to_text
import concurrent.futures
import json

import concurrent.futures
import json

FFMPEG_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "ffmpeg.exe")
FFPROBE_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "ffprobe.exe")
pytesseract.pytesseract.tesseract_cmd = os.path.join(
    os.path.dirname(os.path.abspath(__file__)), "tesseract", "tesseract.exe"
)


VALID_BICS = set()
BIC_TO_BANK_INFO = {}

def load_bic_directory(xml_path: str = "20260417_ED807_full.xml") -> None:
    """
    Загружает справочник БИК из файла ED807.xml
    """
    global VALID_BICS, BIC_TO_BANK_INFO
    
    try:
        tree = ET.parse(xml_path)
        root = tree.getroot()
        
        # Пространство имён XML
        ns = {'ed': 'urn:cbr-ru:ed:v2.0'}
        
        for bic_entry in root.findall('.//ed:BICDirectoryEntry', ns):
            bic = bic_entry.get('BIC')
            if bic:
                VALID_BICS.add(bic)
                
                # Сохраняем информацию о банке
                participant = bic_entry.find('.//ed:ParticipantInfo', ns)
                if participant is not None:
                    bank_name = participant.get('NameP', '')
                    BIC_TO_BANK_INFO[bic] = {
                        'name': bank_name,
                        'region': participant.get('Rgn', ''),
                        'city': participant.get('Nnp', '')
                    }
        
        print(f"Загружено {len(VALID_BICS)} БИК из справочника")
    except Exception as e:
        print(f"Ошибка загрузки справочника БИК: {e}")
        
mp_face_detection = mp.solutions.face_detection
mp_pose = mp.solutions.pose

# ====================================================================
# МОДЕЛИ ДЛЯ ДЕТЕКЦИИ БИОМЕТРИИ (загружаются один раз при старте)
# ====================================================================
from huggingface_hub import hf_hub_download
from ultralytics import YOLO
from transformers import pipeline as hf_pipeline

# YOLOv8s — детекция рукописных подписей (~22 МБ, mAP@50 = 94.5%)
try:
    _sig_model_path = hf_hub_download(
        repo_id="tech4humans/yolov8s-signature-detector",
        filename="yolov8s.pt"
    )
    _sig_model = YOLO(_sig_model_path)
    print("Модель подписей: OK")
except Exception as e:
    _sig_model = None
    print(f"Модель подписей не загружена: {e}")

# SigLIP — zero-shot классификация отпечатков пальцев (~400 МБ)
try:
    _fp_classifier = hf_pipeline(
        "zero-shot-image-classification",
        model="google/siglip-base-patch16-224"
    )
    _fp_labels = ["a fingerprint", "a document with text", "a photo of a person"]
    print("Модель отпечатков: OK")
except Exception as e:
    _fp_classifier = None
    print(f"Модель отпечатков не загружена: {e}")



def is_file_accessible(path: str) -> bool:
    """
    Проверяет реальную возможность чтения файла. Для этого пытаемся получить
    право писать и читать файл
    """
    try:
        # Пытаемся открыть файл на чтение. 
        # Используем бинарный режим 'rb', чтобы не возиться с кодировками.
        with open(path, 'rb'):
            pass
        return True
    except (PermissionError, OSError):
        return False

def forming_table(root_dir: str = ".//") -> pd.DataFrame:
    """
    функция forming_table принимает на вход ссылку на корневую папку (условная база данных) и возвращает
    таблицу с метаданными о каждом файле (название, расширение, путь). также создает дополнительно колонки
    оценки опасности, содержания, категории.
    """

    raw_data = []

    if not os.path.exists(root_dir):
        print(f"Ошибка: такой директории {root_dir} нет на устройстве")
    else:

        for root, dirs, files in os.walk(root_dir): # os.walk возвращает путь до нынешней папки, все папки в наличии, все файлы в наличии

            for name in files:

                try:

                    full_path = os.path.join(root, name) # полный путь = склейка пути до текущей папки и имени файла с расширением

                    if not is_file_accessible(full_path):

                        raw_data.append({
                            "Имя файла": os.path.splitext(name)[0],
                            "Путь": full_path,
                            "Расширение": os.path.splitext(name)[1].lower(),
                            "Дата создания": "НЕТ ДОСТУПА",
                            "Содержание": "НЕТ ДОСТУПА"
                        })
                        continue 

                    ext = os.path.splitext(name)[1].lower() # ext[0] - имя файла, ext[1] - расширение
                    name = os.path.splitext(name)[0]
                    date = os.path.getctime(full_path)
                    date = datetime.datetime.fromtimestamp(date).strftime('%Y-%m-%d %H:%M:%S')
                    
                    file_info = {
                                "Имя файла" : name,
                                "Путь" : full_path,
                                "Расширение": ext,
                                "Дата создания": date
                                }
                    
                    raw_data.append(file_info)
                
                except Exception as e:

                    print(f"Ошбика при обработке метаданных файла {name}: {e}")

                    # ловим другие ошибки (файл удален в процессе поиска и т.д.)

        df = pd.DataFrame(raw_data)
        df['Содержание'] = "NO"
        df['Требуемый УЗ'] = 0.0
        df['Найденные ПДн'] = "NO"
        df['Категории'] = "NO"

    return df

### Служебные функции
def choose_engine(extension: str) -> str:

        extension = extension.lower()
        cases = {
            # Группа 1: PDF (через PyMuPDF/fitz)
            ".pdf": "pdf_engine",

            # Группа 2: Word (через python-docx)
            ".docx": "docx_engine",

            # Группа 3: Текст (через встроенный open)
            ".txt": "text_engine", ".log": "text_engine", ".md": "text_engine", 
            ".xml": "text_engine", ".html": "text_engine", ".htm": "text_engine",
            ".gif": "image_ocr",

            # JSON (встроенный json)
            ".json": "json_engine",

            # RTF (через striprtf)
            ".rtf": "rtf_engine",

            # DOC старый формат
            ".doc": "doc_engine",

            ".xls": "table_engine",

            # Группа 4: Аудио (через Whisper)
            ".mp3": "whisper", ".wav": "whisper", ".m4a": "whisper",
            ".flac": "whisper", ".ogg": "whisper",

            # Группа 5: Изображения (через Tesseract OCR)
            ".png": "image_ocr", ".jpg": "image_ocr", ".jpeg": "image_ocr",
            ".bmp": "image_ocr", ".tiff": "image_ocr", ".tif": "image_ocr",
            ".webp": "image_ocr",

            # Группа 6: Видео (аудио Whisper + кадры Tesseract OCR)
            ".mp4": "video_engine", ".avi": "video_engine",
            ".mkv": "video_engine", ".mov": "video_engine",
            ".webm": "video_engine", ".wmv": "video_engine",

            # Группа 7: Таблицы (pandas - csv + openpyxl - xlsx + xlrd - xls)
            ".csv": "table_engine", ".tsv": "table_engine",
            ".xlsx": "table_engine", ".xls": "table_engine",

            ".parquet" : "binary_engine"
        }
        return cases.get(extension, "skip")

def flatten_json(obj, prefix=""):
                    """Рекурсивно разворачиваем JSON в плоский текст"""
                    parts = []
                    if isinstance(obj, dict):
                        for k, v in obj.items():
                            parts.extend(flatten_json(v, f"{prefix}{k}: "))
                    elif isinstance(obj, list):
                        for item in obj:
                            parts.extend(flatten_json(item, prefix))
                    else:
                        parts.append(f"{prefix}{obj}")
                    return parts

def extract_binary(path: str, min_length: int = 6) -> str:
    """
    Извлекает читаемые ASCII-строки из бинарного файла.
    Нужно для поиска ПДн внутри .doc и других бинарников.
    """
    try:
        with open(path, "rb") as f:
            raw = f.read()
        ascii_strings = re.findall(rb'[ -~]{%d,}' % min_length, raw)
        result = []
        for s in ascii_strings:
            try:
                result.append(s.decode("ascii"))
            except Exception:
                pass
        return "\n".join(result) if result else "БИНАРНИК: ЧИТАЕМЫХ СТРОК НЕ НАЙДЕНО"
    except Exception as e:
        return f"Ошибка бинарника: {e}"

def _detect_signature(path: str) -> bool:
        """
        Детекция подписи через YOLOv8s (обучена на 10k+ образцов подписей).
        Возвращает True если найдена хотя бы одна подпись с confidence > 0.5
        """
        if _sig_model is None:
            return False
        try:
            results = _sig_model(path, verbose=False)
            for r in results:
                for box in r.boxes:
                    if float(box.conf[0]) > 0.5:
                        return True
        except Exception:
            pass
        return False

def _detect_fingerprint(path: str) -> bool:
        """
        Детекция отпечатка пальца через SigLIP zero-shot классификацию.
        Возвращает True если score класса "a fingerprint" > 0.7
        """
        if _fp_classifier is None:
            return False
        try:
            results = _fp_classifier(path, candidate_labels=_fp_labels)
            for r in results:
                if r["label"] == "a fingerprint" and r["score"] > 0.7:
                    return True
        except Exception:
            pass
        return False


def detect_biometry(path: str) -> list:
        """
        Комбинированная детекция биометрии:
        - MediaPipe: лицо, глаза, силуэт тела
        - YOLOv8s: подпись
        - SigLIP: отпечаток пальца
        Возвращает список: ["лицо (2)", "глаза", "подпись"]
        """
        img = cv2.imread(path)
        if img is None:
            return []

        found = []
        rgb = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)

        # Лицо + глаза (MediaPipe)
        try:
            with mp_face_detection.FaceDetection(
                model_selection=1, min_detection_confidence=0.5
            ) as detector:
                results = detector.process(rgb)
                if results.detections:
                    found.append(f"лицо ({len(results.detections)})")
                    for det in results.detections:
                        kp = det.location_data.relative_keypoints
                        if len(kp) >= 2:
                            found.append("глаза")
                            break
        except Exception:
            pass

        # Силуэт тела (MediaPipe Pose)
        try:
            with mp_pose.Pose(
                static_image_mode=True, min_detection_confidence=0.5
            ) as pose:
                if pose.process(rgb).pose_landmarks:
                    found.append("силуэт тела")
        except Exception:
            pass

        # Подпись (YOLOv8s) — принимает path, не gray
        try:
            if _detect_signature(path):
                found.append("подпись")
        except Exception:
            pass

        # Отпечаток пальца (SigLIP) — принимает path, не gray
        try:
            if _detect_fingerprint(path):
                found.append("отпечаток пальца")
        except Exception:
            pass

        return found

### Функция для парсинга одного файла

def worker_parse_file(file_data):
    """
    Автономный обработчик одного файла.
    Принимает: (индекс, путь_к_файлу, расширение)
    Возвращает: (индекс, извлеченный_текст)
    """

    idx, path, ext = file_data
    engine = choose_engine(ext) # функция выбора движка
    
    try:

        # Легкие форматы

        if engine == "json_engine":
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                data = json.load(f)
            # Используем вспомогательную функцию flatten_json (вынеси её тоже в корень)
            lines = flatten_json(data) 
            return idx, "\n".join(str(l) for l in lines).strip() or "ПУСТОЙ JSON"
            

        elif engine == "rtf_engine":
            try:
                with open(path, "r", encoding="utf-8", errors="ignore") as f:
                    raw = f.read()
                # Предполагается наличие функции rtf_to_text
                text = rtf_to_text(raw) 
                return idx, text.strip() or "ПУСТОЙ RTF"
                
            except Exception as e:
                print(f"RTF не прочитан. Детали: {e}")
                

        elif engine == "doc_engine":
            try:
                # Пробуем через antiword (внешняя утилита)
                
                result = subprocess.run(
                    ["antiword", path], 
                    capture_output=True, 
                    text=True, 
                    timeout=30,
                    encoding='utf-8', # явно указываем кодировку для Windows
                    errors='ignore'
                )
                
                
                if result.returncode == 0 and result.stdout.strip():
                    return idx, result.stdout.strip()
                    
                else:
                    # Fallback: если antiword не выдал текст, читаем как бинарник
                    return idx, extract_binary(path)
                    
                    
            except FileNotFoundError:
                # antiword не установлен в системе — используем встроенный бинарный парсер
                return idx, extract_binary(path)
                
            except Exception as e:
                return idx, f"Ошибка DOC (бинарный парсинг): {e}"

        elif engine == "docx_engine":

            try:
                # Открываем документ
                doc = Document(path)
                # Извлекаем текст из всех параграфов
                text = "\n".join([p.text for p in doc.paragraphs])
                
                # Формируем результат
                result_text = text.strip() if text.strip() else "ПУСТОЙ DOCX"
                return idx, result_text
                
            except Exception as e:
                # Возвращаем описание ошибки вместо падения процесса
                return idx, f"Ошибка в чтении файла DOCx: {str(e)}"
        
        elif engine == "text_engine":

            try:

                with open(path, "r", encoding="utf-8", errors="ignore") as txt:
                    content = txt.read().strip()
                
                return idx, content if content else "ПУСТОЙ ТЕКСТОВЫЙ ФАЙЛ"
            except Exception as e:
                return idx, f"Ошибка в чтении текстового формата: {str(e)}"

        elif engine == "table_engine":

            try:
                import pandas as pd
                tbl_ext = os.path.splitext(path)[1].lower()
                tdf = None

                # 1. Чтение таблиц разных форматов
                if tbl_ext == ".csv":
                    for sep in [",", ";", "\t", "|"]:
                        try:
                            tdf = pd.read_csv(path, sep=sep, dtype=str, on_bad_lines="skip", encoding_errors="ignore")
                            if len(tdf.columns) > 1:
                                break
                        except:
                            continue
                    if tdf is None:
                        tdf = pd.read_csv(path, dtype=str, on_bad_lines="skip", encoding_errors="ignore")

                elif tbl_ext == ".tsv":
                    tdf = pd.read_csv(path, sep="\t", dtype=str, on_bad_lines="skip", encoding_errors="ignore")

                elif tbl_ext in (".xlsx", ".xls"):
                    eng = "openpyxl" if tbl_ext == ".xlsx" else "xlrd"
                    # Читаем все листы
                    sheets = pd.read_excel(path, sheet_name=None, dtype=str, engine=eng)
                    parts = []
                    for sname, sdf in sheets.items():
                        # Собираем заголовки и строки листа
                        lines = [" ".join(str(c) for c in sdf.columns)]
                        for _, r in sdf.iterrows():
                            row_t = " ".join(str(v) for v in r.values if pd.notna(v))
                            if row_t.strip():
                                lines.append(row_t)
                        if lines:
                            parts.append(f"[{sname}]\n" + "\n".join(lines))
                    
                    return idx, "\n\n".join(parts) if parts else "ПУСТАЯ ТАБЛИЦА"

                # 2. Обработка CSV/TSV после чтения
                if tdf is not None:
                    lines = [" ".join(str(c) for c in tdf.columns)]
                    for _, r in tdf.iterrows():
                        row_t = " ".join(str(v) for v in r.values if pd.notna(v))
                        if row_t.strip():
                            lines.append(row_t)
                    return idx, "\n".join(lines) if lines else "ПУСТАЯ ТАБЛИЦА"
                
                return idx, "НЕ УДАЛОСЬ ПРОЧИТАТЬ ТАБЛИЦУ"

            except Exception as e:
                return idx, f"Ошибка таблицы: {str(e)}"
        
        ### Тяжеловесные
                

        elif engine == "docx_engine":

            try:
                # Открываем документ
                doc = Document(path)
                # Извлекаем текст из всех параграфов
                text = "\n".join([p.text for p in doc.paragraphs])
                
                # Формируем результат
                result_text = text.strip() if text.strip() else "ПУСТОЙ DOCX"
                return idx, result_text
                
            except Exception as e:
                # Возвращаем описание ошибки вместо падения процесса
                return idx, f"Ошибка в чтении файла DOCx: {str(e)}"
        
        elif engine == "text_engine":

            try:

                with open(path, "r", encoding="utf-8", errors="ignore") as txt:
                    content = txt.read().strip()
                
                return idx, content if content else "ПУСТОЙ ТЕКСТОВЫЙ ФАЙЛ"
            except Exception as e:
                return idx, f"Ошибка в чтении текстового формата: {str(e)}"

        
        ### Тяжеловесные

        elif engine == "pdf_engine":
            try:
                # Порог количества символов.
                # Если на странице больше 500 символов, считаем, что это текстовый документ,
                # и биометрию (фото/подписи) искать не нужно.
                TEXT_THRESHOLD = 300

                with fitz.open(path) as doc:
                    
                    # --- Вспомогательная функция автоповорота ---
                    def fix_rotation(img):
                        try:
                            osd = pytesseract.image_to_osd(img, config='--psm 0', lang='rus+eng', output_type=pytesseract.Output.DICT)
                            angle = osd.get('rotate', 0)
                            if angle != 0:
                                img = img.rotate(-angle, expand=True)
                        except pytesseract.TesseractError:
                            pass
                        return img
                    # -------------------------------------------

                    # --- Основная функция обработки страницы ---
                    def process_page(page):
                        # 1. Пытаемся получить текст из PDF (без OCR)
                        # Используем "textdict" для более точного подсчета, если нужно, 
                        # но get_text() тоже подойдет для оценки объема.
                        raw_text = page.get_text()
                        text_len = len(raw_text.strip())
                        
                        # Если текста очень много, пропускаем тяжелые проверки
                        if text_len > TEXT_THRESHOLD:
                            # Возвращаем текст, биометрия пустая
                            return raw_text.strip(), ""

                        # 2. Если текста мало, это может быть скан или пустая страница.
                        # Делаем рендер, поворот и OCR.
                        pix = page.get_pixmap(matrix=fitz.Matrix(150/72, 150/72))
                        img = Image.open(io.BytesIO(pix.tobytes("png")))
                        
                        # Автоповорот
                        img = fix_rotation(img)
                        
                        # Распознаем текст
                        ocr_text = pytesseract.image_to_string(img, lang="rus+eng")
                        ocr_text_len = len(ocr_text.strip())
                        
                        # Обновляем итоговый текст (берем либо raw, либо OCR, в зависимости от того, что нашли)
                        # Логика: если raw был пустой, берем OCR. Если был небольшой, но OCR нашел больше - берем OCR.
                        final_text = raw_text.strip()
                        if ocr_text_len > text_len:
                            final_text = ocr_text.strip()
                        
                        bio_info = ""
                        
                        # 3. Проверяем биометрию ТОЛЬКО если текста мало.
                        # Если после OCR текста стало много (найден скрытый текст) — тоже пропускаем биометрию.
                        if len(final_text) < TEXT_THRESHOLD:
                            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                                img.save(tmp.name, format="PNG")
                                tmp_path = tmp.name
                            
                            bio = detect_biometry(tmp_path)
                            if bio:
                                bio_info = f"[БИОМЕТРИЯ стр.{page.number + 1}: {', '.join(bio)}]"
                            
                            os.unlink(tmp_path)
                        
                        return final_text, bio_info

                    # Запускаем обработку в параллельных потоках
                    with ThreadPoolExecutor(max_workers=8) as pool:
                        results = list(pool.map(process_page, doc))
                    
                    # Сборка результатов
                    final_text_parts = []
                    bio_results = []
                    
                    for text_part, bio_part in results:
                        if text_part:
                            final_text_parts.append(text_part)
                        if bio_part:
                            bio_results.append(bio_part)

                    full_text = "\n\n".join(final_text_parts)
                    if bio_results:
                        full_text += "\n" + "\n".join(bio_results)

                return idx, full_text if full_text else "ПУСТОЙ ПДФ ФАЙЛ"
            except Exception as e:
                return idx, f"Ошибка в чтении пдф формата: {str(e)}"

        elif engine == "whisper":
            try:
                import whisper
                # Загружаем модель внутри процесса. 
                # "base" весит около 150Мб, это допустимо для параллелизма.
                model = whisper.load_model("base") 

                # Загрузка и обработка аудио
                audio = whisper.load_audio(path)
                audio_segment = whisper.pad_or_trim(audio)

                # Детекция языка
                mel = whisper.log_mel_spectrogram(audio_segment).to(model.device)
                _, probs = model.detect_language(mel)

                detected_language = max(probs, key=probs.get)

                # Транскрибация
                result = model.transcribe(path, language=detected_language)
                
                text = result["text"].strip() if result else ""
                if text:
                    # Добавляем метку биометрии, как и в других воркерах
                    return idx, text + "\n[БИОМЕТРИЯ: образец голоса]"

                else:
                    return idx, "НИЧЕГО НЕ ИЗВЛЕЧЕНО (АУДИО ПУСТОЕ)"

            except Exception as e:
                return idx, f"Произошел сбой при извлечении аудиодорожки: {str(e)}"
                

        elif engine == "image_ocr":
            try:
                from PIL import Image
                # Открываем изображение
                img = Image.open(path)
                
                # Приводим к совместимому формату для Tesseract
                
                # Приводим к совместимому формату для Tesseract
                if img.mode not in ("L", "RGB"):
                    img = img.convert("RGB")
                
                # Извлекаем текст
                text = pytesseract.image_to_string(img, lang="rus+eng").strip()
                
                # Запускаем детекцию лиц, подписей и т.д.
                # Функция detect_biometry должна быть объявлена в корне файла
                # Извлекаем текст
                text = pytesseract.image_to_string(img, lang="rus+eng").strip()
                
                # Запускаем детекцию лиц, подписей и т.д.
                # Функция detect_biometry должна быть объявлена в корне файла
                bio = detect_biometry(path)
                
                
                if bio:
                    text += f"\n[БИОМЕТРИЯ: {', '.join(bio)}]"

                return idx, text if text else "OCR НЕ ИЗВЛЁК ТЕКСТ"
                
            except Exception as e:
                return idx, f"Ошибка OCR: {str(e)}"

        elif engine == "video_engine":
            try:
                import whisper
                import whisper
                results = []
                # Загружаем модель внутри процесса (лучше использовать маленькую для скорости)
                v_model = whisper.load_model("base")

                # 1. Извлекаем аудиодорожку через FFmpeg
                with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as tmp:
                    tmp_audio = tmp.name
                
                subprocess.run([
                    FFMPEG_PATH, "-i", path,
                    "-vn", "-acodec", "pcm_s16le",
                    "-ar", "16000", "-ac", "1", "-y", tmp_audio
                ], capture_output=True, timeout=120)

                # 2. Транскрибация аудио
                if os.path.exists(tmp_audio) and os.path.getsize(tmp_audio) > 1000:
                    audio = whisper.load_audio(tmp_audio)
                    mel = whisper.log_mel_spectrogram(whisper.pad_or_trim(audio)).to(v_model.device)
                    _, probs = v_model.detect_language(mel)
                    mel = whisper.log_mel_spectrogram(whisper.pad_or_trim(audio)).to(v_model.device)
                    _, probs = v_model.detect_language(mel)
                    lang = max(probs, key=probs.get)
                    res = v_model.transcribe(tmp_audio, language=lang)
                    res = v_model.transcribe(tmp_audio, language=lang)
                    if res["text"].strip():
                        results.append(res["text"].strip())
                        results.append("[БИОМЕТРИЯ: образец голоса]")

                # 3. Нарезка кадров и OCR
                with tempfile.TemporaryDirectory() as tmpdir:
                    subprocess.run([
                        FFMPEG_PATH, "-i", path,
                        "-vf", "fps=1/10", "-q:v", "2", "-y",
                        os.path.join(tmpdir, "f_%04d.jpg")
                    ], capture_output=True, timeout=180)

                    prev_text = ""
                    for f_name in sorted(os.listdir(tmpdir)):
                        if not f_name.endswith(".jpg"):
                            continue
                        
                        frame_path = os.path.join(tmpdir, f_name)
                        img = Image.open(frame_path)
                        
                        # OCR кадра
                        t = pytesseract.image_to_string(img, lang="rus+eng").strip()
                        if t and t != prev_text:
                            results.append(t)
                            prev_text = t
                        
                        # Биометрия кадра
                        bio = detect_biometry(frame_path)
                        if bio:
                            results.append(f"[БИОМЕТРИЯ кадр {f_name}: {', '.join(bio)}]")

                return idx, "\n".join(results) if results else "ВИДЕО: ТЕКСТ НЕ ИЗВЛЕЧЁН"

            except Exception as e:
                return idx, f"Ошибка видео: {str(e)}"

        elif engine == "binary_engine":
            try:
                import polars
                # Читаем parquet файл

                temp = polars.read_parquet(path)
                
                # Приводим все колонки к строковому типу и объединяем в текст
                # Используем .sample или .head(1000), если файлы гигантские
                raw_text = " ".join(
                    temp.select(polars.all().cast(polars.Utf8))
                    .to_series()
                    .to_list()
                )
                
                return idx, raw_text if raw_text.strip() else "БИНАРНЫЙ ФАЙЛ ПУСТ"
            
            except Exception as e:
                return idx, f"Ошибка в чтении бинарника (Parquet). Детали: {e}"

        return idx, f"Для файла {ext} был не найден движок."
    except Exception as e:

        print(f"При парсинге произошла ошибка. Детали: {e}")
        return idx, f"При парсинге произошла ошибка. Детали: {e}", 

def parsing(df, update_callback=None):

    """
    
    """
    tasks = [(idx, row["Путь"], row["Расширение"]) for idx, row in df.iterrows()]
    
    with concurrent.futures.ProcessPoolExecutor(max_workers=4) as executor:
        futures = {executor.submit(worker_parse_file, t): t for t in tasks}
        for i, f in enumerate(concurrent.futures.as_completed(futures)):
            res_idx, content = f.result()
            df.at[res_idx, "Содержание"] = content
            if update_callback:
                update_callback("Обработка...", i + 1, len(tasks))
    return df

### Обработка результатов парсинга
### Обработка результатов парсинга
def seek_danger(df: pd.DataFrame, update_callback=None) -> pd.DataFrame:
    """
    Гибридный подход: sentence-embeddings + кластеризация + regex-якоря.
    
    Этап 1: Regex находит «якорные» файлы с железобетонными ПДн (СНИЛС, паспорт, биометрия).
    Этап 2: Все тексты прогоняются через sentence-transformer → эмбеддинги.
    Этап 3: Считаем косинусное сходство каждого файла с якорями.
             Если файл похож на якорный — помечаем как кандидат.
    Этап 4: Для кандидатов запускаем детальный regex для определения типов ПДн.
    """
    import re
    import numpy as np
    from collections import defaultdict
    from sentence_transformers import SentenceTransformer
    from sklearn.metrics.pairwise import cosine_similarity

    # ========================================================================
    # 0. ИНИЦИАЛИЗАЦИЯ МОДЕЛИ (один раз)
    # ========================================================================
    if not hasattr(seek_danger, '_model'):
        print("Загрузка sentence-transformer...")
        # Лёгкая мультиязычная модель, ~120MB, быстрая
        seek_danger._model = SentenceTransformer(
            'sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2'
        )
        print("Модель загружена")

    model = seek_danger._model

    # ========================================================================
    # 1. ВАЛИДАТОРЫ
    # ========================================================================
    VALID_OPERATOR_CODES = {
        "900","901","902","903","904","905","906","908","909",
        "910","911","912","913","914","915","916","917","918","919",
        "920","921","922","923","924","925","926","927","928","929",
        "930","931","932","933","934","936","937","938","939",
        "941","942","949","950","951","952","953","954","955","958","959",
        "960","961","962","963","964","965","966","967","968","969",
        "970","971","977","978","979","980","981","982","983","984",
        "985","986","987","988","989","990","991","992","993","994",
        "995","996","997","999"
    }

    def valid_snils(s):
        d = re.sub(r'\D', '', s)
        if len(d) != 11 or d == "00000000000":
            return False
        t = sum(int(d[i]) * (9 - i) for i in range(9))
        c = t % 101
        if c == 100:
            c = 0
        return c == int(d[9:11])

    def valid_inn12(s):
        s = re.sub(r'\D', '', s)
        if len(s) != 12:
            return False
        c1 = [7, 2, 4, 10, 3, 5, 9, 4, 6, 8]
        t1 = sum(int(s[i]) * c1[i] for i in range(10)) % 11
        if t1 == 10:
            t1 = 0
        if t1 != int(s[10]):
            return False
        c2 = [3, 7, 2, 4, 10, 3, 5, 9, 4, 6, 8]
        t2 = sum(int(s[i]) * c2[i] for i in range(11)) % 11
        if t2 == 10:
            t2 = 0
        return t2 == int(s[11])

    def valid_phone(s):
        d = re.sub(r'\D', '', s)
        if len(d) != 11 or d[0] not in '78':
            return False
        return d[1:4] in VALID_OPERATOR_CODES

    def valid_luhn(s):
        d = re.sub(r'[\s\-]', '', s)
        if not d.isdigit() or len(d) != 16 or d[0] not in '2456':
            return False
        t = 0
        for i, ch in enumerate(d):
            n = int(ch)
            if i % 2 == 0:
                n *= 2
                if n > 9:
                    n -= 9
            t += n
        return t % 10 == 0

    def personal_email(s):
        local = s.split('@')[0].lower()
        bad = {"info", "admin", "support", "noreply", "no-reply", "help",
               "contact", "office", "mail", "sales", "pr", "hr", "secretary",
               "webmaster", "postmaster", "marketing", "service", "feedback",
               "press", "news", "robot", "bot", "auto", "system", "root"}
        return local not in bad and len(local) >= 3

    # ========================================================================
    # 2. ЯКОРНЫЕ ПАТТЕРНЫ (высокая точность, без сомнений = ПДн)
    # ========================================================================
    ANCHOR_PATTERNS = {
        "СНИЛС": (r"\b\d{3}-\d{3}-\d{3}\s?\d{2}\b", valid_snils),
        "Паспорт": (
            r"(?:паспорт|серия)\s*[:\s]*\d{2}\s*\d{2}\s*(?:№|номер)?\s*\d{6}",
            None
        ),
        "Банковская карта": (r"\b(?:\d{4}[\s\-]){3}\d{4}\b", valid_luhn),
        "Биометрия: лицо": (r"\[БИОМЕТРИЯ[^\]]*лицо", None),
        "Биометрия: глаза": (r"\[БИОМЕТРИЯ[^\]]*глаза", None),
        "Биометрия: силуэт": (r"\[БИОМЕТРИЯ[^\]]*силуэт", None),
        "Биометрия: подпись": (r"\[БИОМЕТРИЯ[^\]]*подпись", None),
        "Биометрия: отпечаток": (r"\[БИОМЕТРИЯ[^\]]*отпечаток", None),
        "Биометрия: голос": (r"\[БИОМЕТРИЯ[^\]]*голос", None),
        "MRZ": (r"(?:P<|I<|V<|AC)[A-Z<]{2}[A-Z<]{1,39}<<", None),
        "Медицина": (
            r"(?:диагноз|мкб[\-\s]?\d+|анамнез|эпикриз)\s*[:\-]?\s*[А-Яа-я]",
            None
        ),
        "Согласие_ПДн": (
            r"согласие\s+на\s+обработку\s+персональных",
            None
        ),
    }

    # ========================================================================
    # 3. ДЕТАЛЬНЫЕ ПАТТЕРНЫ (для определения ТИПОВ ПДн в кандидатах)
    # ========================================================================
    DETAIL_PATTERNS = {
        "СНИЛС": (r"\b\d{3}-\d{3}-\d{3}\s?\d{2}\b", valid_snils),
        "Паспорт": (
            r"(?:паспорт|серия)\s*[:\s]*\d{2}\s*\d{2}\s*(?:№|номер)?\s*\d{6}",
            None
        ),
        "Телефон": (
            r"(?:\+7|8)[\s\(-]?\d{3}[\s\)-]?\d{3}[\s-]?\d{2}[\s-]?\d{2}",
            valid_phone
        ),
        "Email": (
            r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,6}",
            personal_email
        ),
        "ИНН": (r"\b\d{12}\b", valid_inn12),
        "Банковская карта": (r"\b(?:\d{4}[\s\-]){3}\d{4}\b", valid_luhn),
        "ФИО": (
            r"[А-ЯЁ][а-яё]{1,30}\s+[А-ЯЁ][а-яё]{1,30}\s+"
            r"[А-ЯЁ][а-яё]{1,30}(?:вич|вна|ич|ична|инична|евич|евна|ович)",
            None
        ),
        "ФИО_инициалы": (
            r"[А-ЯЁ][а-яё]{2,30}\s+[А-ЯЁ]\.\s*[А-ЯЁ]\.",
            None
        ),
        "Дата рождения": (
            r"(?:дата\s+рожд[а-яё]*|родил(?:ся|ась)|д\.?\s*р\.?\s*:|г\.?\s*р\.?)"
            r"\s*[:\-—\s]{0,5}\d{1,2}[.\-/]\d{1,2}[.\-/]\d{2,4}",
            None
        ),
        "Адрес": (
            r"(?:адрес\s+(?:регистрации|проживания|прописки|места\s+жительства)|"
            r"зарегистрирован[а]?\s+по\s+адресу|проживает\s+по\s+адресу)",
            None
        ),
        "Заработная плата": (
            r"(?:зарплата|оклад|заработная\s+плата|з/?п)\s*[:\-]?\s*\d[\d\s]*"
            r"(?:руб|₽|р\.)",
            None
        ),
        "Медицина": (
            r"(?:диагноз|мкб[\-\s]?\d+|анамнез|эпикриз|медицинская\s+карта|"
            r"история\s+болезни)\s*[:\-]?\s*[А-Яа-я]",
            None
        ),
        "Полис ОМС": (
            r"(?:полис\s+омс|омс)\s*[:\-]?\s*\d{16}",
            None
        ),
        "Водительское удостоверение": (
            r"\b[АВЕКМНОРСТУХ]{2}\s?\d{6}\b",
            None
        ),
        "Национальность": (
            r"(?:национальность|национальная\s+принадлежность)\s*[:\-]",
            None
        ),
        "Религиозные убеждения": (
            r"(?:религия|вероисповедание)\s*[:\-]",
            None
        ),
        "Политические убеждения": (
            r"(?:политические\s+убеждения|партийная\s+принадлежность)\s*[:\-]",
            None
        ),
        "Судимость": (
            r"(?:судимость|несудим)\s*[:\-]?\s*(?:есть|нет|не\s+имеется|"
            r"имеется|отсутствует)",
            None
        ),
        "CVV": (r"(?:cvv|cvc|cvv2|cvc2)\s*[:\-]?\s*\d{3}\b", None),
        "Биометрия: лицо": (r"\[БИОМЕТРИЯ[^\]]*лицо", None),
        "Биометрия: глаза": (r"\[БИОМЕТРИЯ[^\]]*глаза", None),
        "Биометрия: силуэт": (r"\[БИОМЕТРИЯ[^\]]*силуэт", None),
        "Биометрия: подпись": (r"\[БИОМЕТРИЯ[^\]]*подпись", None),
        "Биометрия: отпечаток": (r"\[БИОМЕТРИЯ[^\]]*отпечаток", None),
        "Биометрия: голос": (r"\[БИОМЕТРИЯ[^\]]*голос", None),
    }

    # ========================================================================
    # 4. ПОДГОТОВКА ТЕКСТОВ
    # ========================================================================
    total = len(df)
    contents = []
    valid_indices = []     # индексы строк с реальным текстом
    skip_indices = set()   # индексы ошибок/пустых

    for idx, row in df.iterrows():
        content = str(row.get("Содержание", ""))
        if content.startswith("Ошибка") or content.strip() in ("", "nan"):
            skip_indices.add(idx)
            df.at[idx, "Найденные ПДн"] = "Нет никаких нарушений"
        else:
            contents.append(content[:5000])  # первые 5000 символов для эмбеддинга
            valid_indices.append(idx)

    if not valid_indices:
        return df

    print(f"Файлов для анализа: {len(valid_indices)}, пропущено: {len(skip_indices)}")

    # ========================================================================
    # 5. ЭТАП 1: ЯКОРНАЯ РАЗМЕТКА через regex
    # ========================================================================
    anchor_positive = set()   # индексы файлов — точно ПДн
    anchor_negative = set()   # индексы файлов — точно НЕ ПДн
    anchor_found = {}         # idx -> {тип: кол-во}

    CLEAN_INDICATORS = [
        "course syllabus", "curriculum vitae", "conference program",
        "table of contents", "bibliography", "references",
        "abstract:", "isbn", "doi:", "issn:", "proceedings",
        "©", "copyright", "all rights reserved",
        "учебный план", "рабочая программа дисциплины",
        "список литературы", "оглавление", "содержание",
        "тезисы докладов", "сборник трудов",
    ]

    if update_callback:
        update_callback(current_file="Этап 1: якорная разметка",
                       current_file_pos=0, total_files=total)

    for i, idx in enumerate(valid_indices):
        text = str(df.at[idx, "Содержание"])[:200_000]
        text_lower = text.lower()
        found = {}

        for name, (pat, validator) in ANCHOR_PATTERNS.items():
            for m in re.finditer(pat, text_lower, re.IGNORECASE):
                ok = True
                if validator:
                    ok = validator(m.group())
                if ok:
                    found[name] = found.get(name, 0) + 1
                    break

        if found:
            anchor_positive.add(idx)
            anchor_found[idx] = found
        else:
            clean_score = sum(1 for ind in CLEAN_INDICATORS if ind in text_lower)
            if clean_score >= 2 and len(text) > 200:
                anchor_negative.add(idx)

    print(f"Якоря: {len(anchor_positive)} ПДн, {len(anchor_negative)} чистых")

    # ========================================================================
    # 6. ЭТАП 2: ЭМБЕДДИНГИ
    # ========================================================================
    if update_callback:
        update_callback(current_file="Этап 2: вычисление эмбеддингов",
                       current_file_pos=1, total_files=total)

    print("Вычисляем эмбеддинги...")
    embeddings = model.encode(
        contents,
        batch_size=64,
        show_progress_bar=True,
        normalize_embeddings=True   # для cosine similarity через dot product
    )
    print(f"Эмбеддинги: {embeddings.shape}")

    # ========================================================================
    # 7. ЭТАП 3: СКОРИНГ ЧЕРЕЗ СХОДСТВО С ЯКОРЯМИ
    # ========================================================================
    if update_callback:
        update_callback(current_file="Этап 3: скоринг по сходству",
                       current_file_pos=2, total_files=total)

    # Маппинг: позиция в contents[] -> idx в df
    pos_to_idx = {i: idx for i, idx in enumerate(valid_indices)}
    idx_to_pos = {idx: i for i, idx in enumerate(valid_indices)}

    # Собираем эмбеддинги якорных ПДн
    positive_positions = [idx_to_pos[idx] for idx in anchor_positive if idx in idx_to_pos]
    negative_positions = [idx_to_pos[idx] for idx in anchor_negative if idx in idx_to_pos]

    similarity_scores = np.zeros(len(contents))

    if positive_positions:
        pos_embs = embeddings[positive_positions]
        # Средний эмбеддинг "документа с ПДн"
        pos_centroid = pos_embs.mean(axis=0, keepdims=True)

        # Косинусное сходство каждого документа с центроидом ПДн
        sim_to_pd = cosine_similarity(embeddings, pos_centroid).flatten()

        if negative_positions:
            neg_embs = embeddings[negative_positions]
            neg_centroid = neg_embs.mean(axis=0, keepdims=True)
            sim_to_clean = cosine_similarity(embeddings, neg_centroid).flatten()
            # Итоговый скор: разница
            similarity_scores = sim_to_pd - sim_to_clean
        else:
            similarity_scores = sim_to_pd

    print(f"Скоры сходства: min={similarity_scores.min():.3f}, "
          f"max={similarity_scores.max():.3f}, "
          f"mean={similarity_scores.mean():.3f}")

    # ========================================================================
    # 8. ЭТАП 4: ОПРЕДЕЛЕНИЕ КАНДИДАТОВ + ДЕТАЛЬНЫЙ REGEX
    # ========================================================================
    if update_callback:
        update_callback(current_file="Этап 4: детальный анализ кандидатов",
                       current_file_pos=3, total_files=total)

    # Порог сходства: адаптивный
    if positive_positions:
        pos_scores = similarity_scores[positive_positions]
        # Порог = среднее якорных - 1.5 стандартных отклонения
        threshold = pos_scores.mean() - 1.5 * max(pos_scores.std(), 0.05)
        threshold = max(threshold, 0.05)  # не ниже 0.05
    else:
        threshold = 0.1

    print(f"Порог сходства: {threshold:.3f}")

    candidates = set()
    for pos in range(len(contents)):
        idx = pos_to_idx[pos]
        if idx in anchor_positive:
            candidates.add(idx)
        elif similarity_scores[pos] >= threshold:
            candidates.add(idx)

    print(f"Кандидатов для детального анализа: {len(candidates)}")

    # Детальный regex для кандидатов
    for pos, idx in enumerate(valid_indices):
        if update_callback and pos % 100 == 0:
            update_callback(current_file=f"Детализация: {pos}/{len(valid_indices)}",
                           current_file_pos=pos, total_files=len(valid_indices))

        if idx not in candidates:
            df.at[idx, "Найденные ПДн"] = "Нет никаких нарушений"
            continue

        text = str(df.at[idx, "Содержание"])[:200_000]
        text_lower = text.lower()
        pd_count = defaultdict(int)

        # Если уже нашли якорные — добавляем
        if idx in anchor_found:
            for k, v in anchor_found[idx].items():
                if k != "Согласие_ПДн":  # вспомогательный якорь
                    pd_count[k] += v

        # Ищем все детальные паттерны
        for name, (pat, validator) in DETAIL_PATTERNS.items():
            if name in pd_count:
                continue  # уже нашли через якорь

            search_text = text_lower if name in (
                "Паспорт", "Дата рождения", "Адрес", "Заработная плата",
                "Медицина", "Полис ОМС", "Национальность",
                "Религиозные убеждения", "Политические убеждения",
                "Судимость", "CVV"
            ) else text

            for m in re.finditer(pat, search_text, re.IGNORECASE):
                ok = True
                if validator:
                    ok = validator(m.group())
                if ok:
                    real_name = "ФИО" if name == "ФИО_инициалы" else name
                    pd_count[real_name] += 1
                    break

        # --- Финальное решение ---
        if pd_count:
            strong = {"Паспорт", "СНИЛС", "ИНН", "Банковская карта",
                      "Полис ОМС", "Водительское удостоверение", "MRZ",
                      "Дата рождения", "Заработная плата", "Медицина",
                      "Судимость", "Национальность", "Религиозные убеждения",
                      "Политические убеждения", "CVV"}
            bio = {k for k in pd_count if k.startswith("Биометрия")}
            weak = {"ФИО", "Телефон", "Email", "Адрес"}
            weak_found = {t for t in pd_count if t in weak}
            has_strong = any(t in strong for t in pd_count)

            if has_strong or bio or len(weak_found) >= 2:
                parts = [f"{t}({c})" for t, c in pd_count.items()]
                df.at[idx, "Найденные ПДн"] = ",".join(parts)
            else:
                # Слабые сигналы, но файл ОЧЕНЬ похож на якорные
                i_pos = idx_to_pos.get(idx, -1)
                if i_pos >= 0 and similarity_scores[i_pos] >= threshold * 1.5:
                    parts = [f"{t}({c})" for t, c in pd_count.items()]
                    df.at[idx, "Найденные ПДн"] = ",".join(parts)
                else:
                    df.at[idx, "Найденные ПДн"] = "Нет никаких нарушений"
        else:
            # Нет regex-находок, но similarity очень высокий — пометить?
            i_pos = idx_to_pos.get(idx, -1)
            if i_pos >= 0 and similarity_scores[i_pos] >= threshold * 2.0:
                # Очень похож на ПДн, но regex ничего не нашёл
                # Помечаем осторожно — только если парсинг мог пропустить
                df.at[idx, "Найденные ПДн"] = "ПДн(1)"
            else:
                df.at[idx, "Найденные ПДн"] = "Нет никаких нарушений"

    return df



def categories(df: pd.DataFrame) -> pd.DataFrame:
    import re
    from collections import defaultdict

    PD_TO_CATEGORY = {
        "ФИО": "Обычные персональные данные",
        "Телефон": "Обычные персональные данные",
        "Email": "Обычные персональные данные",
        "Дата рождения": "Обычные персональные данные",
        "Адрес": "Обычные персональные данные",
        "Паспорт": "Государственные идентификаторы",
        "СНИЛС": "Государственные идентификаторы",
        "ИНН": "Государственные идентификаторы",
        "Водительское удостоверение": "Государственные идентификаторы",
        "MRZ": "Государственные идентификаторы",
        "Полис ОМС": "Государственные идентификаторы",
        "Банковская карта": "Платежная информация",
        "Банковский счет": "Платежная информация",
        "БИК": "Платежная информация",
        "CVV": "Платежная информация",
        "Биометрия: лицо": "Биометрические данные",
        "Биометрия: глаза": "Биометрические данные",
        "Биометрия: силуэт": "Биометрические данные",
        "Биометрия: подпись": "Биометрические данные",
        "Биометрия: отпечаток": "Биометрические данные",
        "Биометрия: голос": "Биометрические данные",
        "Медицина": "Специальные категории ПДн",
        "Национальность": "Специальные категории ПДн",
        "Раса": "Специальные категории ПДн",
        "Религиозные убеждения": "Специальные категории ПДн",
        "Политические убеждения": "Специальные категории ПДн",
        "Судимость": "Специальные категории ПДн",
        "Заработная плата": "Финансовые данные",
    }

    for idx, row in df.iterrows():
        found_pdns_str = str(row.get("Найденные ПДн", ""))

        if found_pdns_str == "Нет никаких нарушений" or not found_pdns_str:
            df.at[idx, "Категории"] = "Нет нарушений"
            continue

        category_counts = defaultdict(int)

        for item in found_pdns_str.split(","):
            item = item.strip()
            if not item:
                continue
            m = re.match(r'^([A-Za-zА-Яа-яЁё\s:]+)\((\d+)\)$', item)
            if m:
                pd_type = m.group(1).strip()
                count = int(m.group(2))
                cat = PD_TO_CATEGORY.get(pd_type)
                if not cat:
                    for key, val in PD_TO_CATEGORY.items():
                        if key.lower() == pd_type.lower():
                            cat = val
                            break
                if cat:
                    category_counts[cat] += count
                else:
                    category_counts["Неизвестная категория"] += count

        if category_counts:
            parts = [f"{c}({n})" for c, n in sorted(category_counts.items())]
            df.at[idx, "Категории"] = ",".join(parts)
        else:
            df.at[idx, "Категории"] = "Неизвестная категория"

    return df



def evaluate_violations(df: pd.DataFrame) -> pd.DataFrame:
    """
    Функция принимает на вход датафрейм и по колонке "Категории"
    определяет требуемый уровень защищенности (УЗ) от 1 до 4.
    
    Уровни защищенности (Приказ ФСТЭК России №21):
    - УЗ-1: Наличие специальных категорий ПДн или биометрических данных (высокий риск)
    - УЗ-2: Наличие платежной информации или государственных идентификаторов в больших объемах (более 3)
    - УЗ-3: Наличие государственных идентификаторов в небольших объемах (от 1 до 3 включительно) 
            или обычных ПДн в больших объемах (более 3)
    - УЗ-4: Наличие только обычных ПДн в небольших объемах (от 1 до 3 включительно)
    - 0: Нет нарушений
    """
    
    import re
    from collections import defaultdict
    
    # ========================================================================
    # 1. ОПРЕДЕЛЕНИЕ КАТЕГОРИЙ
    # ========================================================================
    
    HIGH_RISK_CATEGORIES = {
        "Специальные категории ПДн",
        "Биометрические данные",
    }
    
    PAYMENT_CATEGORIES = {
        "Платежная информация",
    }
    
    GOV_ID_CATEGORIES = {
        "Государственные идентификаторы",
    }
    
    COMMON_PD_CATEGORIES = {
        "Обычные персональные данные",
        "Финансовые данные",
    }
    
    # ========================================================================
    # 2. ВСПОМОГАТЕЛЬНАЯ ФУНКЦИЯ
    # ========================================================================
    def parse_categories(categories_str: str) -> dict:
        result = defaultdict(int)
        
        if not categories_str or categories_str == "Нет нарушений":
            return result
        
        for item in categories_str.split(","):
            item = item.strip()
            if not item:
                continue
            
            match = re.match(r'^([А-Яа-я\s]+)\((\d+)\)$', item)
            if match:
                category = match.group(1).strip()
                count = int(match.group(2))
                result[category] = count
        
        return result
    
    # ========================================================================
    # 3. ОСНОВНОЙ ЦИКЛ
    # ========================================================================
    for idx, row in df.iterrows():
        categories_str = row["Категории"]
        
        # Если нет нарушений
        if categories_str == "Нет нарушений":
            df.at[idx, "Требуемый УЗ"] = 0 
            continue
        
        categories_counts = parse_categories(categories_str)
        
        if not categories_counts:
            df.at[idx, "Требуемый УЗ"] = 0 
            continue
        
        # ====================================================================
        # УЗ-1: Специальные категории или биометрия
        # ====================================================================
        has_high_risk = False
        for category in HIGH_RISK_CATEGORIES:
            if category in categories_counts and categories_counts[category] > 0:
                has_high_risk = True
                break
        
        if has_high_risk:
            df.at[idx, "Требуемый УЗ"] = 1 
            continue
        
        # ====================================================================
        # УЗ-2: Платежная информация ИЛИ гос. идентификаторы > 3
        # ====================================================================
        has_payment = False
        for category in PAYMENT_CATEGORIES:
            if category in categories_counts and categories_counts[category] > 0:
                has_payment = True
                break
        
        if has_payment:
            df.at[idx, "Требуемый УЗ"] = 2
            continue
        
        gov_id_count = 0
        for category in GOV_ID_CATEGORIES:
            gov_id_count += categories_counts.get(category, 0)
        
        if gov_id_count > 3:
            df.at[idx, "Требуемый УЗ"] = 2
            continue
        
        # ====================================================================
        # УЗ-3: Гос. идентификаторы 1-3 ИЛИ обычные ПДн > 3
        # ====================================================================
        if 1 <= gov_id_count <= 3:
            df.at[idx, "Требуемый УЗ"] = 3
            continue
        
        common_pd_count = 0
        for category in COMMON_PD_CATEGORIES:
            common_pd_count += categories_counts.get(category, 0)
        
        if common_pd_count > 3:
            df.at[idx, "Требуемый УЗ"] = 3
            continue
        
        # ====================================================================
        # УЗ-4: Только обычные ПДн 1-3
        # ====================================================================
        if 1 <= common_pd_count <= 3:
            df.at[idx, "Требуемый УЗ"] = 4
            continue
        
        df.at[idx, "Требуемый УЗ"] = 4
    
    return df

### Главная функция
### Главная функция
def run_scanning(path: str, update_callback = None)->pd.DataFrame:

    """
    Функция для запуска полного сканнирования хранилища.
    Все функции расставлены в порядке их необходимого применения.
    В целом функции независимы, но это рекомендованный порядок

    """
    start_time = time.time()
    root_dir = path

    # --- Шаг 1 ---
    df = forming_table(root_dir=root_dir)
    time_step1 = time.time()
    print(f"Время извлечения метаданных: {round(time_step1 - start_time, 2)} сек.")
    print(df)

    # --- Шаг 2 ---
    extracted_df = parsing(df, update_callback=update_callback) 
    time_step2 = time.time()
    print(f"Время парсинга информации: {round(time_step2 - time_step1, 2)} сек.")
    print(extracted_df)

    # --- Шаг 3 ---
    found_danger_df = seek_danger(extracted_df) 
    time_step3 = time.time()
    print(f"Время анализа по №152-ФЗ: {round(time_step3 - time_step2, 2)} сек.")
    
    # --- Шаг 4 ---
    categorized_df = categories(found_danger_df)
    time_step4 = time.time()
    print(f"Время анализа по №152-ФЗ: {round(time_step4 - time_step3, 2)} сек.")

    # --- Итог ---
    evaluated_df = evaluate_violations(categorized_df)

    time_step5 = time.time() - start_time
    print(f"\nВремя оценки нарушений: {round(time_step4, 2)} сек.")

    evaluated_df.drop(columns=["Содержание"], inplace=True)
    evaluated_df.reset_index(drop=True, inplace=True)

    try:
        conn = sqlite3.connect("DataBase.db")
        evaluated_df.to_sql("scan_results", con = conn, if_exists = "replace")
        print("Успешно создана база данных")
        print(evaluated_df.columns)

    except Exception as e:

        print(f"Создать базу данных не удалось: {e}")

    finally:

        conn.close()

   
    print(evaluated_df["Категории"])
    # Итог
    total_time = time.time() - start_time
    print(f"\nОБЩЕЕ ВРЕМЯ РАБОТЫ: {round(total_time, 2)} сек.")
    
    return evaluated_df